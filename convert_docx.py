# -*- coding: utf-8 -*-
"""
将程序说明文档.txt 转换为 Word 文档，按用户要求设置格式
"""
import re
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_FILE = r"e:\MyApplication\程序说明文档.txt"
OUTPUT_FILE = r"e:\MyApplication\程序说明文档.docx"

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)


def set_run_font(run, font_cn="宋体", font_en="Times New Roman", size=Pt(12)):
    """设置 run 的中英文字体与字号"""
    run.font.size = size
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rpr_fonts = rpr.find(qn("w:rFonts"))
    if rpr_fonts is None:
        rpr_fonts = parse_xml(f"<w:rFonts {nsdecls('w')} />")
        rpr.insert(0, rpr_fonts)
    rpr_fonts.set(qn("w:eastAsia"), font_cn)
    rpr_fonts.set(qn("w:ascii"), font_en)
    rpr_fonts.set(qn("w:hAnsi"), font_en)


def add_paragraph_with_font(text, font_cn="宋体", font_en="Times New Roman",
                            size=Pt(12), bold=False, alignment=None,
                            line_spacing=1.5, first_line_indent=None,
                            space_before=0, space_after=0):
    """添加段落并设置格式"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        p.alignment = alignment
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent

    run = p.add_run(text)
    set_run_font(run, font_cn, font_en, size)
    run.bold = bold
    return p


def set_paragraph_text(paragraph, text, font_cn="宋体", font_en="Times New Roman",
                       size=Pt(12), bold=False, alignment=None,
                       line_spacing=1.5, first_line_indent=None,
                       space_before=0, space_after=0):
    """设置已有段落的文字和格式（清空后重新添加）"""
    paragraph.clear()
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line_indent is not None:
        paragraph.paragraph_format.first_line_indent = first_line_indent

    run = paragraph.add_run(text)
    set_run_font(run, font_cn, font_en, size)
    run.bold = bold


def add_code_block_line(text, size=Pt(10.5)):
    """添加一行代码（五号，1.5行距）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # 算法/代码内容保持缩进
    run = p.add_run(text)
    set_run_font(run, "宋体", "Times New Roman", size)
    return p


def add_separator_line():
    """添加代码框分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("─" * 50)
    set_run_font(run, "宋体", "Times New Roman", Pt(10.5))
    return p


# ── 读取文本 ──
with open(INPUT_FILE, "r", encoding="utf-8") as f:
    lines = f.readlines()

# ── 解析并生成文档 ──
i = 0
in_code_block = False
code_label_text = None  # 记录代码块前的 "算法X.Y" 或 "代码X.Y"

while i < len(lines):
    raw = lines[i]
    stripped = raw.rstrip("\n").rstrip("\r")

    # 分隔线判断（代码块边界）
    if re.match(r"^─{10,}$", stripped):
        if not in_code_block:
            # 开始代码块：先输出算法/代码标签（如果有）
            if code_label_text:
                add_paragraph_with_font(
                    code_label_text, bold=False, size=Pt(10.5),
                    line_spacing=1.5, space_before=6, space_after=0
                )
                code_label_text = None
            add_separator_line()
            in_code_block = True
        else:
            add_separator_line()
            in_code_block = False
        i += 1
        continue

    if in_code_block:
        # 代码块内的行直接输出
        add_code_block_line(stripped)
        i += 1
        continue

    # ── 标题识别 ──
    # # 标题 —— 文档大标题（仅第一行）
    if stripped.startswith("# ") and not stripped.startswith("## "):
        add_paragraph_with_font(
            stripped[2:], bold=True, size=Pt(16),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.5, space_before=12, space_after=12
        )
        i += 1
        continue

    if stripped.startswith("## ") and not stripped.startswith("### "):
        # 章标题（如: ## 1 绪论）
        text = stripped[3:]
        add_paragraph_with_font(
            text, bold=True, size=Pt(14),
            line_spacing=1.5, space_before=12, space_after=6
        )
        i += 1
        continue

    if stripped.startswith("### ") and not stripped.startswith("#### "):
        # 节标题（如: ### 1.1 项目背景）
        text = stripped[4:]
        add_paragraph_with_font(
            text, bold=True, size=Pt(12),
            line_spacing=1.5, space_before=8, space_after=4
        )
        i += 1
        continue

    if stripped.startswith("#### ") and not stripped.startswith("##### "):
        # 条标题（如: #### 1.1.1 开发工具与环境）
        text = stripped[5:]
        add_paragraph_with_font(
            text, bold=True, size=Pt(12),
            line_spacing=1.5, space_before=6, space_after=3
        )
        i += 1
        continue

    if stripped.startswith("##### "):
        # 更低级别标题
        text = stripped[6:]
        add_paragraph_with_font(
            text, bold=True, size=Pt(12),
            line_spacing=1.5, space_before=4, space_after=2
        )
        i += 1
        continue

    # ── 空行 ──
    if stripped == "":
        i += 1
        continue

    # ── 算法/代码标签识别 ──
    # 匹配 "算法 X.Y ..." 或 "代码 X.Y ..."
    m = re.match(r"^(算法|代码)\s+\d+\.\d+\s+.*$", stripped)
    if m:
        # 如果下一行是分隔线，则暂存为代码块标签
        next_line = lines[i + 1].rstrip("\n").rstrip("\r") if i + 1 < len(lines) else ""
        if re.match(r"^─{10,}$", next_line):
            code_label_text = stripped
            i += 1
            continue
        else:
            # 否则作为普通正文
            pass

    # ── 正文段落 ──
    # 正文格式：宋体/Times New Roman, 小四(12pt), 1.5行距, 首行缩进2字符
    add_paragraph_with_font(
        stripped,
        first_line_indent=Pt(24),  # 小四约12pt，2字符约24pt
        line_spacing=1.5,
        space_before=0,
        space_after=0
    )
    i += 1

# ── 保存 ──
doc.save(OUTPUT_FILE)
print(f"文档已生成: {OUTPUT_FILE}")
